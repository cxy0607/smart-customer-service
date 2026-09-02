"""文档加载模块：把 PDF / Word 文件解析为 LangChain Document

设计要点（设计说明）：
- 按扩展名分派到对应解析器，新增文件类型只需加一个分支
- PDF 逐页提取并记录页码到 metadata，回答时能展示「答案出自第几页」，提升可信度
- 解析失败抛出业务异常，由上层记录到文档状态机（failed + error_msg），不拖垮整个请求
"""
from pathlib import Path

from langchain_core.documents import Document
from pypdf import PdfReader

from app.core.exceptions import BusinessError, ErrorCode

# 支持的文件类型
SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


def load_pdf(file_path: Path) -> list[Document]:
    """解析 PDF：逐页提取文本，metadata 记录来源文件名与页码"""
    docs: list[Document] = []
    try:
        reader = PdfReader(str(file_path))
        for page_no, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            text = text.strip()
            if not text:
                continue  # 跳过空白页（如纯图片页）
            docs.append(
                Document(
                    page_content=text,
                    metadata={"source": file_path.name, "page": page_no},
                )
            )
    except Exception as e:
        raise BusinessError(ErrorCode.RAG_ERROR, f"PDF 解析失败: {e}") from e
    return docs


def load_docx(file_path: Path) -> list[Document]:
    """解析 Word：按段落提取文本"""
    try:
        import docx  # 延迟导入，避免非 Word 场景加载开销

        document = docx.Document(str(file_path))
        paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
        if not paragraphs:
            raise BusinessError(ErrorCode.RAG_ERROR, "Word 文档中没有可提取的文本内容")
        return [
            Document(
                page_content="\n".join(paragraphs),
                metadata={"source": file_path.name, "page": 1},
            )
        ]
    except BusinessError:
        raise
    except Exception as e:
        raise BusinessError(ErrorCode.RAG_ERROR, f"Word 解析失败: {e}") from e


def load_document(file_path: Path) -> list[Document]:
    """统一入口：按扩展名分派解析器"""
    ext = file_path.suffix.lower()
    if ext == ".pdf":
        return load_pdf(file_path)
    if ext == ".docx":
        return load_docx(file_path)
    raise BusinessError(ErrorCode.PARAM_ERROR, f"不支持的文件类型: {ext}，仅支持 PDF / Word")
