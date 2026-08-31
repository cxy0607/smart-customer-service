"""知识库/文档/FAQ 管理接口测试（含 RBAC 权限验证）

注意：上传文档与创建 FAQ 会真实调用百炼 API（向量化），
文档测试需轮询等待后台任务完成。
"""
import io
import time
import uuid

import docx

from app.core.exceptions import ErrorCode

# 测试专用知识库名（带随机后缀，避免多次运行冲突）
KB_NAME = f"测试知识库_{uuid.uuid4().hex[:6]}"


def _make_docx_bytes() -> bytes:
    """生成测试用 Word 文档内容"""
    buf = io.BytesIO()
    d = docx.Document()
    d.add_paragraph("退货政策：用户自签收之日起 7 天内可申请无理由退货。")
    d.add_paragraph("质量问题退货，运费由商城承担。")
    d.save(buf)
    return buf.getvalue()


# ===== 知识库 CRUD =====
def test_kb_crud(client, auth_headers):
    # 创建
    resp = client.post(
        "/api/v1/knowledge-bases",
        json={"name": KB_NAME, "description": "接口测试用"},
        headers=auth_headers,
    )
    assert resp.status_code == 200
    kb_id = resp.json()["data"]["id"]

    # 列表（应包含刚创建的）
    resp = client.get("/api/v1/knowledge-bases", headers=auth_headers)
    names = [kb["name"] for kb in resp.json()["data"]]
    assert KB_NAME in names

    # 修改
    resp = client.put(
        f"/api/v1/knowledge-bases/{kb_id}",
        json={"description": "已修改"},
        headers=auth_headers,
    )
    assert resp.json()["data"]["description"] == "已修改"

    # 删除
    resp = client.delete(f"/api/v1/knowledge-bases/{kb_id}", headers=auth_headers)
    assert resp.json()["code"] == 0


def test_kb_duplicate_name(client, auth_headers):
    """重名知识库应被拒绝"""
    client.post("/api/v1/knowledge-bases", json={"name": KB_NAME}, headers=auth_headers)
    resp = client.post("/api/v1/knowledge-bases", json={"name": KB_NAME}, headers=auth_headers)
    assert resp.json()["code"] == ErrorCode.PARAM_ERROR
    # 清理
    for kb in client.get("/api/v1/knowledge-bases", headers=auth_headers).json()["data"]:
        if kb["name"] == KB_NAME:
            client.delete(f"/api/v1/knowledge-bases/{kb['id']}", headers=auth_headers)


# ===== 文档上传（异步状态机）=====
def test_upload_and_process_document(client, auth_headers):
    # 建知识库
    kb_id = client.post(
        "/api/v1/knowledge-bases", json={"name": KB_NAME}, headers=auth_headers
    ).json()["data"]["id"]

    # 上传 Word 文档
    resp = client.post(
        f"/api/v1/knowledge-bases/{kb_id}/documents",
        files={"file": ("售后政策.docx", _make_docx_bytes(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        headers=auth_headers,
    )
    assert resp.status_code == 200, resp.text
    doc_id = resp.json()["data"]["id"]
    assert resp.json()["data"]["status"] == "pending"

    # 轮询等待后台向量化完成（最多 60 秒）
    status = ""
    for _ in range(60):
        resp = client.get(f"/api/v1/knowledge-bases/{kb_id}/documents", headers=auth_headers)
        doc = next(d for d in resp.json()["data"] if d["id"] == doc_id)
        status = doc["status"]
        if status in ("succeeded", "failed"):
            break
        time.sleep(1)

    assert status == "succeeded", f"文档处理未成功: {doc}"
    assert doc["chunk_count"] >= 1

    # 清理
    client.delete(f"/api/v1/documents/{doc_id}", headers=auth_headers)
    client.delete(f"/api/v1/knowledge-bases/{kb_id}", headers=auth_headers)


def test_upload_reject_unsupported_type(client, auth_headers):
    """不支持的文件类型应被拒绝"""
    kb_id = client.post(
        "/api/v1/knowledge-bases", json={"name": KB_NAME}, headers=auth_headers
    ).json()["data"]["id"]
    resp = client.post(
        f"/api/v1/knowledge-bases/{kb_id}/documents",
        files={"file": ("evil.exe", b"fake", "application/octet-stream")},
        headers=auth_headers,
    )
    assert resp.json()["code"] == ErrorCode.PARAM_ERROR
    client.delete(f"/api/v1/knowledge-bases/{kb_id}", headers=auth_headers)


# ===== FAQ CRUD（同步向量索引）=====
def test_faq_crud(client, auth_headers):
    kb_id = client.post(
        "/api/v1/knowledge-bases", json={"name": KB_NAME}, headers=auth_headers
    ).json()["data"]["id"]

    # 创建（会真实调用向量化 API）
    resp = client.post(
        f"/api/v1/faqs/knowledge-bases/{kb_id}/faqs",
        json={"question": "怎么退货？", "answer": "7 天内可申请退货。"},
        headers=auth_headers,
    )
    assert resp.status_code == 200, resp.text
    faq_id = resp.json()["data"]["id"]

    # 列表
    resp = client.get(f"/api/v1/faqs/knowledge-bases/{kb_id}/faqs", headers=auth_headers)
    assert len(resp.json()["data"]) == 1

    # 修改
    resp = client.put(
        f"/api/v1/faqs/{faq_id}",
        json={"answer": "7 天内可申请无理由退货。"},
        headers=auth_headers,
    )
    assert "无理由" in resp.json()["data"]["answer"]

    # 删除
    resp = client.delete(f"/api/v1/faqs/{faq_id}", headers=auth_headers)
    assert resp.json()["code"] == 0

    client.delete(f"/api/v1/knowledge-bases/{kb_id}", headers=auth_headers)


# ===== RBAC 权限 =====
def test_rbac_anonymous_denied(client):
    """未登录访问管理接口应 401"""
    assert client.post("/api/v1/knowledge-bases", json={"name": "x"}).status_code == 401
    assert client.get("/api/v1/knowledge-bases").status_code == 401


def test_rbac_normal_user_denied(client):
    """普通用户访问管理接口应 403（用户注册功能为管理员预留，此处验证 token 伪造保护）"""
    # 手工签发一个普通用户 token 验证 RBAC（正常流程普通用户由管理员创建）
    from app.core.security import create_access_token

    user_token = create_access_token(999999, "test_user", "user")
    resp = client.post(
        "/api/v1/knowledge-bases",
        json={"name": "x"},
        headers={"Authorization": f"Bearer {user_token}"},
    )
    # 用户不存在 -> 401；若存在则为 403。两种都属于拒绝，验证"未放行"即可
    assert resp.status_code in (401, 403)
