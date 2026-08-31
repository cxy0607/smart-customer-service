"""对话模块测试：FAQ 优先命中、RAG 流式生成、SSE 协议、会话管理、限流

注意：FAQ/文档向量化会真实调用百炼 API，测试耗时数秒属正常
"""
import io
import json
import time
import uuid

import docx

from app.core.exceptions import ErrorCode

def _make_docx_bytes() -> bytes:
    """生成含退货政策内容的测试 Word 文档"""
    buf = io.BytesIO()
    d = docx.Document()
    d.add_paragraph("退货政策：用户自签收商品之日起 7 天内，商品保持完好可申请无理由退货。")
    d.add_paragraph("若因商品质量问题退货，运费由商城承担。")
    d.save(buf)
    return buf.getvalue()


def _parse_sse(body: str) -> list[dict]:
    """解析 SSE 响应体为事件列表 [{event, data}, ...]"""
    events = []
    for block in body.split("\n\n"):
        ev = {}
        for line in block.strip().split("\n"):
            if line.startswith("event: "):
                ev["event"] = line[len("event: "):]
            elif line.startswith("data: "):
                ev["data"] = json.loads(line[len("data: "):])
        if ev:
            events.append(ev)
    return events


def _create_kb(client, headers) -> int:
    """创建测试知识库（每次随机名，避免测试间重名冲突）"""
    resp = client.post(
        "/api/v1/knowledge-bases",
        json={"name": f"对话测试库_{uuid.uuid4().hex[:8]}"},
        headers=headers,
    )
    assert resp.status_code == 200, f"知识库创建失败: {resp.text}"
    return resp.json()["data"]["id"]


def _wait_document_ready(client, headers, kb_id, doc_id, timeout=60):
    """轮询等待文档向量化完成，返回最终状态"""
    for _ in range(timeout):
        resp = client.get(f"/api/v1/knowledge-bases/{kb_id}/documents", headers=headers)
        doc = next(d for d in resp.json()["data"] if d["id"] == doc_id)
        if doc["status"] in ("succeeded", "failed"):
            return doc
        time.sleep(1)
    raise TimeoutError("文档处理超时")


# ===== 核心流程：FAQ 命中 =====
def test_chat_faq_hit(client, auth_headers):
    """用户提问与 FAQ 相似时，直接返回预设答案（不走大模型）"""
    kb_id = _create_kb(client, auth_headers)
    # 创建 FAQ（问题向量化）
    resp = client.post(
        f"/api/v1/faqs/knowledge-bases/{kb_id}/faqs",
        json={"question": "怎么退货？", "answer": "签收后 7 天内可申请无理由退货。"},
        headers=auth_headers,
    )
    assert resp.status_code == 200, resp.text
    faq_id = resp.json()["data"]["id"]

    # 用相近问法提问
    resp = client.post(
        "/api/v1/chat",
        json={"kb_id": kb_id, "message": "如何办理退货呢"},
        headers=auth_headers,
    )
    assert resp.status_code == 200
    events = _parse_sse(resp.text)
    kinds = [e["event"] for e in events]
    assert kinds == ["meta", "delta", "done"], f"事件序列异常: {kinds}"

    meta = events[0]["data"]
    assert meta["match_type"] == "faq"
    assert "7 天内可申请无理由退货" in events[-1]["data"]["answer"]

    # 清理
    client.delete(f"/api/v1/faqs/{faq_id}", headers=auth_headers)
    client.delete(f"/api/v1/knowledge-bases/{kb_id}", headers=auth_headers)


# ===== 核心流程：RAG 流式生成 =====
def test_chat_rag_stream(client, auth_headers):
    """未命中 FAQ 时走 RAG：检索知识库文档 + LLM 流式生成"""
    kb_id = _create_kb(client, auth_headers)

    # 上传文档并等待处理完成
    resp = client.post(
        f"/api/v1/knowledge-bases/{kb_id}/documents",
        files={"file": ("售后政策.docx", _make_docx_bytes(), "application/octet-stream")},
        headers=auth_headers,
    )
    doc_id = resp.json()["data"]["id"]
    doc = _wait_document_ready(client, auth_headers, kb_id, doc_id)
    assert doc["status"] == "succeeded", f"文档处理失败: {doc}"

    # 提问（知识库内容，不命中 FAQ）
    resp = client.post(
        "/api/v1/chat",
        json={"kb_id": kb_id, "message": "退货的运费谁来承担？"},
        headers=auth_headers,
    )
    assert resp.status_code == 200
    events = _parse_sse(resp.text)

    meta = events[0]["data"]
    assert meta["match_type"] == "rag"
    # 引用来源应包含文件名
    assert any("售后政策" in s["source"] for s in meta["sources"])

    # 流式增量拼接后与 done 的完整答案一致
    deltas = "".join(e["data"]["text"] for e in events if e["event"] == "delta")
    done = [e for e in events if e["event"] == "done"][0]
    assert deltas == done["data"]["answer"]
    assert "商城" in done["data"]["answer"]  # 基于知识库的正确答案

    # 消息落库检查：用户 1 条 + 助手 1 条，助手消息带引用来源
    conv_id = meta["conversation_id"]
    resp = client.get(f"/api/v1/conversations/{conv_id}/messages", headers=auth_headers)
    messages = resp.json()["data"]
    assert len(messages) == 2
    assert messages[0]["role"] == "user"
    assert messages[1]["role"] == "assistant"
    assert messages[1]["source_docs"]  # 引用来源已持久化

    # 清理
    client.delete(f"/api/v1/documents/{doc_id}", headers=auth_headers)
    client.delete(f"/api/v1/knowledge-bases/{kb_id}", headers=auth_headers)


# ===== 会话管理 =====
def test_conversation_list_and_delete(client, auth_headers):
    kb_id = _create_kb(client, auth_headers)

    # 第一次提问会创建会话
    resp = client.post(
        "/api/v1/chat",
        json={"kb_id": kb_id, "message": "你好"},
        headers=auth_headers,
    )
    conv_id = _parse_sse(resp.text)[0]["data"]["conversation_id"]

    # 会话列表包含该会话，标题取问题前 30 字
    resp = client.get("/api/v1/conversations", headers=auth_headers)
    convs = resp.json()["data"]
    assert any(c["id"] == conv_id and c["title"] == "你好" for c in convs)

    # 删除会话
    resp = client.delete(f"/api/v1/conversations/{conv_id}", headers=auth_headers)
    assert resp.json()["code"] == 0
    resp = client.get("/api/v1/conversations", headers=auth_headers)
    assert all(c["id"] != conv_id for c in resp.json()["data"])

    client.delete(f"/api/v1/knowledge-bases/{kb_id}", headers=auth_headers)


# ===== 限流 =====
def test_rate_limit_sliding_window():
    """滑动窗口限流单元测试：limit=2 时第 3 次请求被拒绝"""
    from app.core.rate_limit import check_rate_limit

    key = f"unit-test:{uuid.uuid4().hex}"
    assert check_rate_limit(key, limit=2) is True
    assert check_rate_limit(key, limit=2) is True
    assert check_rate_limit(key, limit=2) is False  # 超限拒绝
